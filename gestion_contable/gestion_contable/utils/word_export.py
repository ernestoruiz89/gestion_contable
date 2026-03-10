import hashlib
import io
import re
from html import unescape

import frappe
from frappe import _
from frappe.utils import cint, cstr, now_datetime

from gestion_contable.gestion_contable.utils.estados_financieros import get_package_column_labels

REPORT_TITLE = "Informe Completo de EEFF Auditados"
NON_AUDITED_REPORT_TITLE = "Estados Financieros y Notas"
REMITTANCE_TITLE = "Carta de Remision"
FONT_NAME = "Arial Narrow"
BODY_SIZE = 12
COMPACT_SIZE = 10


def export_financial_package_to_word(package_name):
    package = frappe.get_doc("Paquete Estados Financieros Cliente", package_name)

    document = _build_financial_package_document(package)
    stream = io.BytesIO()
    document.save(stream)
    content = stream.getvalue()
    stream.close()

    from frappe.utils.file_manager import save_file

    file_name = _build_non_audited_word_export_filename(package.name)
    file_doc = save_file(file_name, content, "Paquete Estados Financieros Cliente", package.name, is_private=1)
    version_info = _register_package_document_version(
        package.name,
        tipo_documento="Word Revision Cliente",
        file_doc=file_doc,
        content_hash=hashlib.sha256(content).hexdigest(),
        estado_documento="Generado",
        observaciones="Documento Word EEFF generado para revision del cliente.",
    )
    return {
        "file_name": file_doc.file_name,
        "file_url": file_doc.file_url,
        "file_id": file_doc.name,
        "version_documento": version_info["version_documento"],
        "tipo_documento": version_info["tipo_documento"],
    }


def export_audited_financial_package_to_word(package_name):
    package = frappe.get_doc("Paquete Estados Financieros Cliente", package_name)
    if not package.dictamen_de_auditoria:
        frappe.throw(
            _("El paquete debe tener un Dictamen de Auditoria vinculado para exportarse a Word."),
            title=_("Dictamen Requerido"),
        )

    dictamen = frappe.get_doc("Informe Final Auditoria", package.dictamen_de_auditoria)
    if dictamen.tipo_de_informe != "Dictamen de Auditoria":
        frappe.throw(
            _("El documento vinculado en Dictamen de Auditoria no corresponde a un dictamen valido."),
            title=_("Dictamen Invalido"),
        )

    document = _build_audited_package_document(package, dictamen)
    stream = io.BytesIO()
    document.save(stream)
    content = stream.getvalue()
    stream.close()

    from frappe.utils.file_manager import save_file

    file_name = _build_word_export_filename(package.name)
    file_doc = save_file(file_name, content, "Paquete Estados Financieros Cliente", package.name, is_private=1)
    version_info = _register_package_document_version(
        package.name,
        tipo_documento="Word Revision Cliente",
        file_doc=file_doc,
        content_hash=hashlib.sha256(content).hexdigest(),
        estado_documento="Generado",
        observaciones="Documento Word generado para revision del cliente.",
    )
    return {
        "file_name": file_doc.file_name,
        "file_url": file_doc.file_url,
        "file_id": file_doc.name,
        "version_documento": version_info["version_documento"],
        "tipo_documento": version_info["tipo_documento"],
    }


def export_remittance_letter_to_word(package_name):
    package = frappe.get_doc("Paquete Estados Financieros Cliente", package_name)
    if not package.dictamen_de_auditoria:
        frappe.throw(
            _("El paquete debe tener un Dictamen de Auditoria vinculado para exportar la carta de remision."),
            title=_("Dictamen Requerido"),
        )

    dictamen = frappe.get_doc("Informe Final Auditoria", package.dictamen_de_auditoria)
    if dictamen.tipo_de_informe != "Dictamen de Auditoria":
        frappe.throw(
            _("El documento vinculado en Dictamen de Auditoria no corresponde a un dictamen valido."),
            title=_("Dictamen Invalido"),
        )

    document = _build_remittance_document(package, dictamen)
    stream = io.BytesIO()
    document.save(stream)
    content = stream.getvalue()
    stream.close()

    from frappe.utils.file_manager import save_file

    file_name = _build_remittance_filename(package.name)
    file_doc = save_file(file_name, content, "Paquete Estados Financieros Cliente", package.name, is_private=1)
    return {
        "file_name": file_doc.file_name,
        "file_url": file_doc.file_url,
        "file_id": file_doc.name,
    }


def _build_audited_package_document(package, dictamen):
    Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, OxmlElement, qn, Cm, Pt, RGBColor = _docx_imports()

    document = Document()
    _configure_document(document, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, Pt, document_title=REPORT_TITLE)

    title = document.add_paragraph(REPORT_TITLE, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(package.razon_social_reportante or package.cliente or "Cliente")
    _set_paragraph_runs_font(subtitle)

    badge = document.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge.add_run("Documento formal de revision del cliente").italic = True
    _set_paragraph_runs_font(badge)

    cover_table = document.add_table(rows=4, cols=2)
    cover_pairs = [
        ("Cliente", package.cliente or "-"),
        ("Razon social reportante", package.razon_social_reportante or "-"),
        ("Identificacion fiscal", package.identificacion_fiscal_reportante or "-"),
        ("Periodo", package.periodo_contable or "-"),
        ("Fecha de corte", cstr(package.fecha_corte or "-")),
        ("Fecha de emision", cstr(package.fecha_emision or "-")),
        ("Marco contable", package.marco_contable or "-"),
        ("Version", cstr(package.version or "-")),
    ]
    for index, pair in enumerate(cover_pairs):
        row = cover_table.rows[index // 2]
        cell = row.cells[index % 2]
        _fill_label_value_cell(cell, pair[0], pair[1])
    _style_meta_table(cover_table)

    summary = document.add_paragraph(
        "Se presenta el juego completo de estados financieros auditados, sus notas y el dictamen de auditoria independiente para revision del cliente."
    )
    _set_paragraph_runs_font(summary)

    document.add_page_break()
    toc_heading = document.add_paragraph("Indice", style="Heading 1")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_paragraph = document.add_paragraph()
    _append_field(toc_paragraph, 'TOC \\o "1-3" \\h \\z \\u', "Actualice la tabla de contenido al abrir el documento en Word.", OxmlElement, qn)
    note = document.add_paragraph("Nota: si Word no actualiza el indice automaticamente, use Actualizar campo.")
    _set_paragraph_runs_font(note)

    document.add_page_break()
    _add_dictamen_section(document, dictamen)
    document.add_page_break()
    _add_estados_section(document, package)
    document.add_page_break()
    _add_notas_section(document, package)
    return document


def _build_financial_package_document(package):
    Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, OxmlElement, qn, Cm, Pt, RGBColor = _docx_imports()

    document = Document()
    _configure_document(document, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, Pt, document_title=NON_AUDITED_REPORT_TITLE)

    title = document.add_paragraph(NON_AUDITED_REPORT_TITLE, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(package.razon_social_reportante or package.cliente or "Cliente")
    _set_paragraph_runs_font(subtitle)

    badge = document.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge.add_run("Documento preparado por la firma para revision del cliente").italic = True
    _set_paragraph_runs_font(badge)

    cover_table = document.add_table(rows=4, cols=2)
    cover_pairs = [
        ("Cliente", package.cliente or "-"),
        ("Razon social reportante", package.razon_social_reportante or "-"),
        ("Identificacion fiscal", package.identificacion_fiscal_reportante or "-"),
        ("Periodo", package.periodo_contable or "-"),
        ("Fecha de corte", cstr(package.fecha_corte or "-")),
        ("Fecha de emision", cstr(package.fecha_emision or "-")),
        ("Marco contable", package.marco_contable or "-"),
        ("Version", cstr(package.version or "-")),
    ]
    for index, pair in enumerate(cover_pairs):
        row = cover_table.rows[index // 2]
        cell = row.cells[index % 2]
        _fill_label_value_cell(cell, pair[0], pair[1])
    _style_meta_table(cover_table)

    summary = document.add_paragraph(
        "Se presenta el juego de estados financieros y sus notas explicativas preparados por la firma para revision del cliente."
    )
    _set_paragraph_runs_font(summary)

    document.add_page_break()
    toc_heading = document.add_paragraph("Indice", style="Heading 1")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_paragraph = document.add_paragraph()
    _append_field(toc_paragraph, 'TOC \\o "1-3" \\h \\z \\u', "Actualice la tabla de contenido al abrir el documento en Word.", OxmlElement, qn)
    note = document.add_paragraph("Nota: si Word no actualiza el indice automaticamente, use Actualizar campo.")
    _set_paragraph_runs_font(note)

    document.add_page_break()
    _add_estados_section(document, package)
    document.add_page_break()
    _add_notas_section(document, package)
    return document


def _build_remittance_document(package, dictamen):
    Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, OxmlElement, qn, Cm, Pt, RGBColor = _docx_imports()

    document = Document()
    _configure_document(document, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, Pt, document_title=REPORT_TITLE)
    _add_delivery_section(document, package, dictamen)
    return document


def _docx_imports():
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError:
        frappe.throw(
            _("No esta instalada la dependencia opcional python-docx. Instala la dependencia en el sitio con <b>bench pip install python-docx</b> para habilitar la exportacion Word."),
            title=_("Dependencia Faltante"),
        )
    return Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, OxmlElement, qn, Cm, Pt, RGBColor


def _configure_document(document, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, Pt, document_title=REPORT_TITLE):
    styles = document.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"].font.size = Pt(BODY_SIZE)
    styles["Title"].font.name = FONT_NAME
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.name = FONT_NAME
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = FONT_NAME
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.name = FONT_NAME
    styles["Heading 3"].font.size = Pt(BODY_SIZE)

    for section in document.sections:
        _configure_section(section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=False, document_title=document_title)


def _configure_section(section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=False, document_title=REPORT_TITLE):
    _Document, _Align, WD_ORIENTATION, _SectionStart, _OxmlElement, _qn, _Cm, _Pt, _RGBColor = _docx_imports()
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.2)
    section.orientation = WD_ORIENTATION.LANDSCAPE if landscape else WD_ORIENTATION.PORTRAIT
    if landscape:
        section.page_width, section.page_height = section.page_height, section.page_width

    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.text = f"{package.razon_social_reportante or package.cliente or '-'} | {document_title}"
    _set_paragraph_runs_font(header_para, size=10)

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = ""
    footer_para.add_run(f"{package.periodo_contable or '-'} | ")
    _append_field(footer_para, "PAGE", "1", OxmlElement, qn)
    footer_para.add_run(" de ")
    _append_field(footer_para, "NUMPAGES", "1", OxmlElement, qn)
    footer_para.add_run(f" | Generado {now_datetime().strftime('%Y-%m-%d %H:%M')}")
    _set_paragraph_runs_font(footer_para, size=10)


def _add_oriented_section(document, package, landscape=False, document_title=REPORT_TITLE):
    _Document, WD_ALIGN_PARAGRAPH, _WD_ORIENTATION, WD_SECTION_START, OxmlElement, qn, Cm, Pt, _RGBColor = _docx_imports()
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    _configure_section(section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=landscape, document_title=document_title)
    return section


def _append_field(paragraph, instruction, placeholder, OxmlElement, qn):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def _add_dictamen_section(document, dictamen):
    opinion_standard = "NIA 700" if dictamen.tipo_opinion == "Favorable" else "NIA 705"
    document.add_paragraph("Dictamen de Auditoria", style="Heading 1")
    meta = document.add_paragraph()
    meta.add_run(f"Tipo de opinion: {dictamen.tipo_opinion or '-'}")
    meta.add_run(f" | Destinatario: {dictamen.destinatario or '-'}")
    meta.add_run(f" | Fecha: {dictamen.fecha_informe or '-'}")
    _set_paragraph_runs_font(meta)

    _add_rich_section(document, f"Opinion del Auditor ({opinion_standard})", dictamen.opinion_o_conclusion)
    _add_rich_section(document, f"Fundamento de la Opinion ({opinion_standard})", dictamen.fundamento_opinion)

    if dictamen.tipo_opinion in ("Con Salvedades", "Adversa", "Abstencion"):
        _add_rich_section(document, "Asunto que Origina la Modificacion (NIA 705)", dictamen.asunto_que_origina_modificacion)
        modified_title = "Fundamento de la Abstencion (NIA 705)" if dictamen.tipo_opinion == "Abstencion" else "Fundamento de la Opinion Modificada (NIA 705)"
        _add_rich_section(document, modified_title, dictamen.fundamento_salvedad)

    if dictamen.parrafo_enfasis:
        _add_rich_section(document, "Parrafo de Enfasis (NIA 706)", dictamen.parrafo_enfasis)
    if dictamen.parrafo_otros_asuntos:
        _add_rich_section(document, "Parrafo de Otros Asuntos (NIA 706)", dictamen.parrafo_otros_asuntos)

    _add_rich_section(document, "Responsabilidades de la Administracion", dictamen.responsabilidades_administracion)
    _add_rich_section(document, "Responsabilidades del Auditor", dictamen.responsabilidades_auditor)
    _add_rich_section(document, "Conclusion Final", dictamen.conclusion_final)

    sign = document.add_paragraph()
    sign.add_run(dictamen.firmado_por or "-").bold = True
    if dictamen.cargo_firmante:
        sign.add_run(f"\n{dictamen.cargo_firmante}")
    _set_paragraph_runs_font(sign)


def _add_estados_section(document, package):
    document.add_paragraph("Estados Financieros", style="Heading 1")
    estados = frappe.get_all(
        "Estado Financiero Cliente",
        filters={"paquete_estados_financieros_cliente": package.name},
        fields=["name", "tipo_estado", "titulo_formal", "subtitulo", "orden_presentacion"],
        order_by="orden_presentacion asc, creation asc",
        limit_page_length=100,
    )
    if not estados:
        paragraph = document.add_paragraph("No hay estados financieros registrados para este paquete.")
        _set_paragraph_runs_font(paragraph)
        return

    for index, estado in enumerate(estados):
        if index > 0:
            document.add_page_break()
        estado_doc = frappe.get_doc("Estado Financiero Cliente", estado.name)
        labels = get_package_column_labels(
            package,
            fecha_actual=estado_doc.fecha_corte or package.fecha_corte,
            fecha_comparativa=estado_doc.fecha_comparativa or getattr(package, "fecha_corte_comparativa", None),
        )
        document.add_paragraph(estado_doc.titulo_formal or estado_doc.tipo_estado, style="Heading 2")
        if estado_doc.subtitulo:
            subtitle = document.add_paragraph(estado_doc.subtitulo)
            _set_paragraph_runs_font(subtitle)
        table = document.add_table(rows=1, cols=4)
        headers = ["Descripcion", "Nota", labels["actual"], labels["comparativo"]]
        for header_index, title in enumerate(headers):
            table.rows[0].cells[header_index].text = title
        total_rows = []
        subtotal_rows = []
        for linea in estado_doc.lineas:
            if getattr(linea, "no_imprimir", 0):
                continue
            line_index = len(table.rows)
            row = table.add_row().cells
            row[0].text = ("    " * max((linea.nivel or 1) - 1, 0)) + (linea.descripcion or "-")
            row[1].text = cstr(linea.numero_nota_referencial or "-")
            row[2].text = _fmt_number(linea.monto_actual)
            row[3].text = _fmt_number(linea.monto_comparativo)
            if linea.es_total:
                total_rows.append(line_index)
            elif linea.es_subtotal:
                subtotal_rows.append(line_index)
        _style_financial_table(table, total_rows=total_rows, subtotal_rows=subtotal_rows)


def _add_notas_section(document, package):
    document.add_paragraph("Notas a los Estados Financieros", style="Heading 1")
    labels = get_package_column_labels(package)
    notas = frappe.get_all(
        "Nota Estado Financiero",
        filters={"paquete_estados_financieros_cliente": package.name},
        fields=["name", "numero_nota", "titulo", "orden_presentacion"],
        order_by="orden_presentacion asc, creation asc",
        limit_page_length=300,
    )
    if not notas:
        paragraph = document.add_paragraph("No hay notas registradas para este paquete.")
        _set_paragraph_runs_font(paragraph)
        return

    for note_index, nota in enumerate(notas):
        if note_index > 0:
            document.add_page_break()
        nota_doc = frappe.get_doc("Nota Estado Financiero", nota.name)
        structured_sections = nota_doc.get_structured_sections()
        document.add_paragraph(f"Nota {nota_doc.numero_nota}. {nota_doc.titulo or 'Sin titulo'}", style="Heading 2")

        if nota_doc.contenido_narrativo:
            _add_rich_block(document, nota_doc.contenido_narrativo)

        if nota_doc.cifras_nota:
            table = document.add_table(rows=1, cols=3)
            headers = ["Concepto", labels["actual"], labels["comparativo"]]
            for header_index, title in enumerate(headers):
                table.rows[0].cells[header_index].text = title
            total_actual = 0.0
            total_comparativo = 0.0
            for cifra in nota_doc.cifras_nota:
                row = table.add_row().cells
                concept_parts = [cstr(cifra.concepto or "-")]
                if cifra.subconcepto:
                    concept_parts.append(cstr(cifra.subconcepto))
                if cifra.comentario:
                    concept_parts.append(cstr(cifra.comentario))
                row[0].text = "\n".join([part for part in concept_parts if part])
                row[1].text = _fmt_number(cifra.monto_actual)
                row[2].text = _fmt_number(cifra.monto_comparativo)
                total_actual += float(cifra.monto_actual or 0)
                total_comparativo += float(cifra.monto_comparativo or 0)
            total_row_index = len(table.rows)
            total_row = table.add_row().cells
            total_row[0].text = "Total"
            total_row[1].text = _fmt_number(total_actual)
            total_row[2].text = _fmt_number(total_comparativo)
            _style_note_table(table, total_rows=[total_row_index])

        for section in structured_sections:
            if section.get("titulo_seccion"):
                document.add_paragraph(section.get("titulo_seccion"), style="Heading 3")
            compact = len(section.get("columnas") or []) > 5
            if section.get("tipo_seccion") in ("Tabla", "Texto y Tabla"):
                if compact:
                    _add_oriented_section(document, package, landscape=True, document_title=REPORT_TITLE)
                    if section.get("titulo_seccion"):
                        document.add_paragraph(section.get("titulo_seccion"), style="Heading 3")
                table = _add_structured_note_table(document, section)
                total_rows = []
                subtotal_rows = []
                for row_index, fila in enumerate(section.get("filas") or [], start=len(table.rows)):
                    row = table.add_row().cells
                    row[0].text = ("    " * max((fila.get("nivel") or 1) - 1, 0)) + cstr(fila.get("descripcion") or "-")
                    for idx, celda in enumerate(fila.get("celdas") or [], start=1):
                        if idx >= len(row):
                            break
                        if celda.get("valor_texto"):
                            row[idx].text = cstr(celda.get("valor_texto"))
                        elif celda.get("valor_numero") is not None:
                            row[idx].text = _fmt_number(celda.get("valor_numero"))
                        else:
                            row[idx].text = "-"
                    tipo_fila = cstr(fila.get("tipo_fila") or "Detalle")
                    if tipo_fila == "Total":
                        total_rows.append(row_index)
                    elif tipo_fila == "Subtotal":
                        subtotal_rows.append(row_index)
                _style_note_table(table, compact=compact, total_rows=total_rows, subtotal_rows=subtotal_rows)
            if section.get("contenido_narrativo"):
                _add_rich_block(document, section.get("contenido_narrativo"))
            if compact:
                _add_oriented_section(document, package, landscape=False, document_title=REPORT_TITLE)

        if nota_doc.observaciones_preparacion:
            document.add_paragraph("Observaciones", style="Heading 3")
            paragraph = document.add_paragraph(cstr(nota_doc.observaciones_preparacion))
            _set_paragraph_runs_font(paragraph)
        document.add_paragraph()


def _add_structured_note_table(document, section):
    columnas = section.get("columnas") or []
    table = document.add_table(rows=1, cols=max(len(columnas) + 1, 1))
    if section.get("tiene_grupos_columnas"):
        first = table.rows[0].cells
        first[0].text = "Concepto"
        col_index = 1
        for grupo in section.get("grupos_columnas") or []:
            for columna in grupo.get("columns") or []:
                first[col_index].text = cstr(columna.get("etiqueta") or "-")
                col_index += 1
        group_row = table.add_row().cells
        group_row[0].text = ""
        col_index = 1
        for grupo in section.get("grupos_columnas") or []:
            label = grupo.get("label") or ""
            for _ in grupo.get("columns") or []:
                group_row[col_index].text = label
                col_index += 1
    else:
        header = table.rows[0].cells
        header[0].text = "Concepto"
        for idx, columna in enumerate(columnas, start=1):
            header[idx].text = cstr(columna.get("etiqueta") or "-")
    return table


def _add_delivery_section(document, package, dictamen):
    document.add_paragraph("Carta de Remision", style="Heading 1")

    fecha = document.add_paragraph(cstr(package.fecha_emision or "-"))
    _set_paragraph_runs_font(fecha)

    destinatario = document.add_paragraph()
    destinatario.add_run(cstr(dictamen.destinatario or package.razon_social_reportante or package.cliente or "-")).bold = True
    _set_paragraph_runs_font(destinatario)

    saludo = document.add_paragraph("Estimados senores:")
    _set_paragraph_runs_font(saludo)

    cuerpo = [
        "Por medio de la presente remitimos el Informe Completo de EEFF Auditados correspondiente al periodo indicado, integrado por el dictamen de auditoria independiente, los estados financieros y sus notas explicativas.",
        f"El documento se entrega para su revision formal y para los fines que estimen convenientes. Identificacion del paquete: {package.name}.",
        "Agradeceremos canalizar cualquier observacion o comentario a traves de los medios previamente acordados con la firma.",
    ]
    for bloque in cuerpo:
        paragraph = document.add_paragraph(bloque)
        _set_paragraph_runs_font(paragraph)

    despedida = document.add_paragraph("Atentamente,")
    _set_paragraph_runs_font(despedida)

    firmante = document.add_paragraph()
    firmante.add_run(dictamen.firmado_por or "____________________________").bold = True
    if dictamen.cargo_firmante:
        firmante.add_run(f"\n{dictamen.cargo_firmante}")
    _set_paragraph_runs_font(firmante)

    referencia = document.add_paragraph()
    referencia.add_run("Referencia: ").bold = True
    referencia.add_run(f"Paquete {package.name} | Dictamen {package.dictamen_de_auditoria or '-'}")
    _set_paragraph_runs_font(referencia)


def _fill_label_value_cell(cell, label, value):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.add_run(f"{label}\n").bold = True
    paragraph.add_run(cstr(value or "-"))
    _set_paragraph_runs_font(paragraph)


def _set_paragraph_runs_font(paragraph, font_name=FONT_NAME, size=BODY_SIZE, bold=False):
    _Document, _Align, _Orientation, _SectionStart, OxmlElement, qn, _Cm, Pt, RGBColor = _docx_imports()
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        if bold:
            run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:eastAsia"), font_name)


def _style_table_no_borders(table):
    _Document, _Align, _Orientation, _SectionStart, OxmlElement, qn, _Cm, _Pt, _RGBColor = _docx_imports()
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "nil")


def _style_cell_border(cell, top=None, bottom=None):
    _Document, _Align, _Orientation, _SectionStart, OxmlElement, qn, _Cm, _Pt, _RGBColor = _docx_imports()
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, spec in (("top", top), ("bottom", bottom)):
        if not spec:
            continue
        edge_el = tcBorders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tcBorders.append(edge_el)
        edge_el.set(qn("w:val"), spec.get("val", "single"))
        edge_el.set(qn("w:sz"), str(spec.get("sz", 8)))
        edge_el.set(qn("w:color"), spec.get("color", "000000"))


def _style_financial_table(table, compact=False, total_rows=None, subtotal_rows=None):
    total_rows = total_rows or []
    subtotal_rows = subtotal_rows or []
    _style_table_no_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                _set_paragraph_runs_font(
                    paragraph,
                    size=COMPACT_SIZE if compact else BODY_SIZE,
                    bold=(row_index == 0 or row_index in total_rows or row_index in subtotal_rows),
                )
                if cell_index > 0:
                    paragraph.alignment = 2
        if row_index in subtotal_rows:
            for cell in row.cells:
                _style_cell_border(cell, top={"val": "single", "sz": 8}, bottom={"val": "single", "sz": 8})
        if row_index in total_rows:
            for cell in row.cells:
                _style_cell_border(cell, top={"val": "single", "sz": 8}, bottom={"val": "double", "sz": 12})


def _style_note_table(table, compact=False, total_rows=None, subtotal_rows=None):
    _style_financial_table(table, compact=compact, total_rows=total_rows, subtotal_rows=subtotal_rows)


def _style_meta_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _set_paragraph_runs_font(paragraph)


def _add_rich_block(document, raw_value):
    text = _html_to_text(raw_value)
    if not text:
        return
    for block in [item.strip() for item in re.split(r"\n\s*\n", text) if item.strip()]:
        paragraph = document.add_paragraph(block)
        _set_paragraph_runs_font(paragraph)


def _add_rich_section(document, heading, raw_value, level=2):
    document.add_paragraph(heading, style=f"Heading {level}")
    text = _html_to_text(raw_value)
    if not text:
        paragraph = document.add_paragraph("Sin contenido.")
        _set_paragraph_runs_font(paragraph)
        return
    for block in [item.strip() for item in re.split(r"\n\s*\n", text) if item.strip()]:
        paragraph = document.add_paragraph(block)
        _set_paragraph_runs_font(paragraph)


def _html_to_text(value):
    text = cstr(value or "")
    if not text.strip():
        return ""
    replacements = (
        (r"<br\s*/?>", "\n"),
        (r"</p>", "\n\n"),
        (r"</div>", "\n"),
        (r"</li>", "\n"),
        (r"<li[^>]*>", "- "),
        (r"</h[1-6]>", "\n"),
    )
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _fmt_number(value):
    if value in (None, ""):
        return "-"
    try:
        return f"{float(value):,.2f}"
    except Exception:
        return cstr(value)


def _build_remittance_filename(package_name):
    safe_name = _sanitize_filename(package_name)
    base_name = f"{REMITTANCE_TITLE} - {safe_name}" if safe_name else REMITTANCE_TITLE
    max_base_length = 135 - len(".docx")
    base_name = base_name[:max_base_length].rstrip(" -_")
    return f"{base_name}.docx"


def _build_word_export_filename(package_name):
    safe_name = _sanitize_filename(package_name)
    base_name = f"{REPORT_TITLE} - {safe_name}" if safe_name else REPORT_TITLE
    max_base_length = 135 - len(".docx")
    base_name = base_name[:max_base_length].rstrip(" -_")
    return f"{base_name}.docx"


def _build_non_audited_word_export_filename(package_name):
    safe_name = _sanitize_filename(package_name)
    base_name = f"{NON_AUDITED_REPORT_TITLE} - {safe_name}" if safe_name else NON_AUDITED_REPORT_TITLE
    max_base_length = 135 - len(".docx")
    base_name = base_name[:max_base_length].rstrip(" -_")
    return f"{base_name}.docx"


def _sanitize_filename(value):
    filename = re.sub(r'[\\/:*?"<>|]+', '-', cstr(value or REPORT_TITLE)).strip()
    return filename or REPORT_TITLE


def _register_package_document_version(package_name, tipo_documento, file_doc, content_hash=None, estado_documento="Generado", observaciones=None):
    package = frappe.get_doc("Paquete Estados Financieros Cliente", package_name)
    rows = list(package.get("versiones_documento_eeff") or [])
    current_versions = [cint(row.version_documento) for row in rows if row.tipo_documento == tipo_documento]
    next_version = max(current_versions or [0]) + 1

    for row in rows:
        if row.tipo_documento == tipo_documento and cint(row.es_version_vigente):
            row.es_version_vigente = 0
            if row.estado_documento == "Generado":
                row.estado_documento = "Reemplazado"

    package.append("versiones_documento_eeff", {
        "tipo_documento": tipo_documento,
        "version_documento": next_version,
        "estado_documento": estado_documento,
        "es_version_vigente": 1,
        "fecha_generacion": now_datetime(),
        "generado_por": frappe.session.user,
        "archivo_file": file_doc.name,
        "nombre_archivo": file_doc.file_name,
        "archivo_url": file_doc.file_url,
        "hash_sha256": content_hash,
        "observaciones": observaciones,
    })
    package.save(ignore_permissions=True)
    return {
        "tipo_documento": tipo_documento,
        "version_documento": next_version,
    }
